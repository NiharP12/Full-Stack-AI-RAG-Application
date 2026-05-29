"""
File processor — extracts text from supported file types.
Supports PDF, TXT, DOCX, CSV, XLSX with intelligent text conversion.
"""

import csv
import io
import logging
import os
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)


class FileProcessor:
    """Extract text content from various file formats."""

    def process(self, file_path: str) -> Dict:
        """
        Process a file and extract its text.

        Args:
            file_path: Absolute path to the file.

        Returns:
            Dict with ``text`` (str), ``pages`` (list of page dicts for PDFs),
            ``file_type`` (str), and ``metadata`` (dict).
        """
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)

        logger.info("Processing file: %s (type: %s, size: %d bytes)", filename, ext, file_size)

        processors = {
            ".pdf": self._process_pdf,
            ".txt": self._process_txt,
            ".docx": self._process_docx,
            ".doc": self._process_docx,  # Try docx parser for .doc as best-effort
            ".csv": self._process_csv,
            ".xlsx": self._process_xlsx,
        }

        processor = processors.get(ext)
        if not processor:
            raise ValueError(f"Unsupported file type: {ext}")

        result = processor(file_path)
        result["file_type"] = ext
        result["metadata"] = {
            "source": filename,
            "file_type": ext,
            "file_size": file_size,
        }

        logger.info("Extracted %d characters from %s", len(result.get("text", "")), filename)
        return result

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def _process_pdf(self, file_path: str) -> Dict:
        """Extract text from PDF using PyPDF2 with pdfplumber fallback."""
        pages = []
        full_text = ""

        # Try PyPDF2 first
        try:
            import PyPDF2

            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    text = text.strip()
                    if text:
                        pages.append({"text": text, "page": i + 1})
                        full_text += text + "\n\n"

            if full_text.strip():
                logger.info("PDF extracted with PyPDF2: %d pages", len(pages))
                return {"text": full_text.strip(), "pages": pages}

        except Exception as exc:
            logger.warning("PyPDF2 extraction failed, trying pdfplumber: %s", exc)

        # Fallback to pdfplumber
        try:
            import pdfplumber

            pages = []
            full_text = ""
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    text = text.strip()
                    if text:
                        pages.append({"text": text, "page": i + 1})
                        full_text += text + "\n\n"

            logger.info("PDF extracted with pdfplumber: %d pages", len(pages))
            return {"text": full_text.strip(), "pages": pages}

        except Exception as exc:
            logger.error("All PDF extraction methods failed: %s", exc)
            raise ValueError(f"Could not extract text from PDF: {exc}")

    # ------------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------------

    def _process_txt(self, file_path: str) -> Dict:
        """Read plain text file."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    text = f.read()
                return {"text": text.strip(), "pages": []}
            except (UnicodeDecodeError, UnicodeError):
                continue

        raise ValueError("Could not decode text file with any supported encoding")

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def _process_docx(self, file_path: str) -> Dict:
        """Extract text from DOCX using python-docx."""
        try:
            import docx

            doc = docx.Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n\n".join(paragraphs)
            return {"text": full_text, "pages": []}

        except Exception as exc:
            raise ValueError(f"Could not extract text from DOCX: {exc}")

    # ------------------------------------------------------------------
    # CSV
    # ------------------------------------------------------------------

    def _process_csv(self, file_path: str) -> Dict:
        """
        Convert CSV rows into natural-language text for better semantic search.

        Example output per row:
            "Product: Laptop, Sales: 500, Month: January"
        """
        try:
            import pandas as pd

            df = pd.read_csv(file_path, encoding="utf-8", on_bad_lines="skip")
            return self._dataframe_to_text(df)

        except Exception as exc:
            # Fallback: basic CSV reader
            try:
                return self._csv_basic_read(file_path)
            except Exception:
                raise ValueError(f"Could not process CSV: {exc}")

    def _csv_basic_read(self, file_path: str) -> Dict:
        """Fallback CSV reader using the csv module."""
        rows_text = []
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.DictReader(f)
            for i, row in enumerate(reader):
                parts = [f"{k}: {v}" for k, v in row.items() if v and v.strip()]
                if parts:
                    rows_text.append(", ".join(parts))
                if i >= 5000:  # Safety limit
                    break

        return {"text": "\n".join(rows_text), "pages": []}

    # ------------------------------------------------------------------
    # XLSX
    # ------------------------------------------------------------------

    def _process_xlsx(self, file_path: str) -> Dict:
        """Convert XLSX rows into natural-language text."""
        try:
            import pandas as pd

            # Read all sheets
            xls = pd.ExcelFile(file_path, engine="openpyxl")
            all_text = []

            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                sheet_result = self._dataframe_to_text(df, sheet_name=sheet_name)
                all_text.append(sheet_result["text"])

            return {"text": "\n\n".join(all_text), "pages": []}

        except Exception as exc:
            raise ValueError(f"Could not process XLSX: {exc}")

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _dataframe_to_text(self, df, sheet_name: Optional[str] = None) -> Dict:
        """
        Convert a pandas DataFrame into natural-language rows.

        Each row becomes: "Column1: value1, Column2: value2, ..."
        """
        import pandas as pd

        rows_text = []

        if sheet_name:
            rows_text.append(f"--- Sheet: {sheet_name} ---")

        # Add column summary
        columns_str = ", ".join(str(c) for c in df.columns)
        rows_text.append(f"Columns: {columns_str}")
        rows_text.append(f"Total rows: {len(df)}")
        rows_text.append("")

        # Convert each row to natural language
        for idx, row in df.iterrows():
            parts = []
            for col in df.columns:
                val = row[col]
                if pd.notna(val) and str(val).strip():
                    parts.append(f"{col}: {val}")
            if parts:
                rows_text.append(", ".join(parts))

            if idx >= 5000:  # Safety limit
                rows_text.append(f"... (truncated, {len(df) - 5000} more rows)")
                break

        return {"text": "\n".join(rows_text), "pages": []}
